# app/services/extractor.py

import os
import re
import pdfplumber
from docx import Document
from fastapi import HTTPException
from app.core.logging_config import logger


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts plain text from a PDF file.
    Logs progress and any issues found.
    """

    logger.info(f"Starting PDF extraction: {os.path.basename(file_path)}")

    extracted_text = []
    pages_processed = 0
    pages_skipped = 0

    try:
        with pdfplumber.open(file_path) as pdf:

            total_pages = len(pdf.pages)
            logger.debug(f"PDF has {total_pages} page(s)")

            for page_number, page in enumerate(pdf.pages, start=1):

                page_text = page.extract_text()

                if page_text:
                    extracted_text.append(page_text)
                    pages_processed += 1
                else:
                    logger.warning(
                        f"Page {page_number} has no extractable text "
                        f"(might be image-based)"
                    )
                    pages_skipped += 1

    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail={
                "error_code": "PDF_EXTRACTION_FAILED",
                "message": f"Failed to extract text from PDF.",
                "hint": "Make sure the PDF is not corrupted or "
                        "password protected."
            }
        )

    logger.info(
        f"PDF extraction complete — "
        f"{pages_processed} pages extracted, "
        f"{pages_skipped} pages skipped"
    )

    full_text = "\n".join(extracted_text)
    return full_text


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts plain text from a DOCX file.
    Logs progress and any issues found.
    """

    logger.info(
        f"Starting DOCX extraction: {os.path.basename(file_path)}"
    )

    extracted_text = []

    try:
        document = Document(file_path)

        total_paragraphs = len(document.paragraphs)
        logger.debug(f"DOCX has {total_paragraphs} paragraph(s)")

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text
            if paragraph_text.strip():
                extracted_text.append(paragraph_text)

    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail={
                "error_code": "DOCX_EXTRACTION_FAILED",
                "message": "Failed to extract text from DOCX.",
                "hint": "Make sure the file is a valid .docx file, "
                        "not an older .doc format."
            }
        )

    non_empty = len(extracted_text)
    logger.info(
        f"DOCX extraction complete — "
        f"{non_empty} non-empty paragraphs extracted"
    )

    full_text = "\n".join(extracted_text)
    return full_text


def clean_text(raw_text: str) -> str:
    """
    Cleans and normalizes raw extracted text.
    """

    logger.debug("Cleaning extracted text...")

    # Replace tabs with spaces
    text = raw_text.replace("\t", " ")

    # Replace multiple spaces with single space
    text = re.sub(r" +", " ", text)

    # Strip each line
    lines = text.split("\n")
    lines = [line.strip() for line in lines]
    text = "\n".join(lines)

    # Replace 3+ blank lines with 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip entire text
    text = text.strip()

    logger.debug(
        f"Text cleaned — "
        f"{len(text)} characters, "
        f"{len(text.split())} words"
    )

    return text


def extract_text(file_path: str) -> str:
    """
    Main extraction function.
    Detects file type and calls the right extractor.
    """

    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    logger.info(f"Extracting text from {file_extension.upper()} file")

    # Choose the right extractor
    if file_extension == ".pdf":
        raw_text = extract_text_from_pdf(file_path)

    elif file_extension == ".docx":
        raw_text = extract_text_from_docx(file_path)

    else:
        logger.error(f"Unsupported file type: {file_extension}")
        raise HTTPException(
            status_code=400,
            detail={
                "error_code": "UNSUPPORTED_FILE_TYPE",
                "message": f"File type {file_extension} is not supported.",
                "hint": "Please upload a PDF or DOCX file."
            }
        )

    # Clean the text
    clean = clean_text(raw_text)

    # Check if we actually got text
    if not clean:
        logger.error(
            "Text extraction returned empty result. "
            "File may be image-based."
        )
        raise HTTPException(
            status_code=422,
            detail={
                "error_code": "NO_TEXT_EXTRACTED",
                "message": "Could not extract any text from the file.",
                "hint": "The file might be a scanned/image-based PDF. "
                        "Please upload a text-based resume created "
                        "digitally in Word or Google Docs."
            }
        )

    logger.info(
        f"Extraction successful — "
        f"{len(clean.split())} words, "
        f"{len(clean)} characters"
    )

    return clean